import os
import pypandoc
import shutil

from docx import Document
from dotenv import load_dotenv

# Load .env variables
load_dotenv()

# Settings
input_dir = os.getenv("INPUT_DIR")
output_dir = os.getenv("OUTPUT_DIR")
find_text = os.getenv("FIND_TEXT")
replace_text = os.getenv("REPLACE_TEXT")
log_file = os.path.join(output_dir, "log.txt")

os.makedirs(output_dir, exist_ok=True)


def log(message):
    with open(log_file, "a") as f:
        f.write(message + "\n")
    print(message)


def convert_doc_to_docx(doc_path, docx_path):
    try:
        pypandoc.convert_file(
            doc_path, "docx", outputfile=docx_path, extra_args=["--standalone"])
        log(f"[CONVERTED] {doc_path} -> {docx_path}")
        return docx_path
    except Exception as e:
        log(f"[ERROR] Failed to convert {doc_path}: {e}")
        return None


def replace_text_in_docx(docx_path, output_path):
    try:
        doc = Document(docx_path)
        replaced = False
        for para in doc.paragraphs:
            if find_text in para.text:
                para.text = para.text.replace(find_text, replace_text)
                replaced = True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find_text in cell.text:
                        cell.text = cell.text.replace(find_text, replace_text)
                        replaced = True
        doc.save(output_path)
        if replaced:
            log(f"[MODIFIED] {docx_path} -> {output_path}")
        else:
            log(f"[COPIED] {docx_path} (no changes)")
        return True
    except Exception as e:
        log(f"[ERROR] Failed to process {docx_path}: {e}")
        return False


# Clear log file at start
with open(log_file, "w") as f:
    f.write("=== Processing Log ===\n")

# Main loop
for filename in os.listdir(input_dir):
    input_path = os.path.join(input_dir, filename)

    try:
        if filename.lower().endswith(".doc"):
            # Convert to docx first
            temp_docx = os.path.join(output_dir, filename + ".docx")
            converted = convert_doc_to_docx(input_path, temp_docx)
            if converted:
                replace_text_in_docx(converted, temp_docx)

        elif filename.lower().endswith(".docx"):
            output_path = os.path.join(output_dir, filename)
            replace_text_in_docx(input_path, output_path)

        else:
            # Copy non-doc files unchanged
            shutil.copy(input_path, output_dir)
            log(f"[SKIPPED] {filename} (not .doc/.docx)")

    except Exception as e:
        log(f"[ERROR] Unexpected error on {filename}: {e}")

log("=== Done! ===")
